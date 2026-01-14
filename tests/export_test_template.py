import os
import sys
import logging
from abc import ABC, abstractmethod
from pathlib import Path

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(levelname)s: %(message)s', stream=sys.stdout)
logger = logging.getLogger(__name__)

class ExportTestRunner(ABC):
    """
    Base class for testing export extensions.
    """
    def __init__(self, input_md_path, output_dir):
        self.input_path = Path(input_md_path)
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.html_content = ""
        self.output_file = None

    def load_and_render(self):
        """Loads Markdown and renders to HTML using the core renderer."""
        if not self.input_path.exists():
            raise FileNotFoundError(f"Input file not found: {self.input_path}")
            
        logger.info(f"Reading {self.input_path}...")
        with open(self.input_path, "r", encoding="utf-8") as f:
            md_text = f.read()
            
        # Dynamic import to avoid circular checks or path issues until run time
        from docnexus.core.renderer import render_baseline
        logger.info("Rendering Markdown to HTML...")
        self.html_content, _ = render_baseline(md_text)

    @abstractmethod
    def run_export(self):
        """Execute the specific export logic. Must set self.output_file."""
        pass

    def verify_file_exists(self):
        if not self.output_file or not self.output_file.exists():
            logger.error(f"Output file not found: {self.output_file}")
            return False
        if self.output_file.stat().st_size == 0:
            logger.error(f"Output file is empty: {self.output_file}")
            return False
        logger.info(f"Export Success: {self.output_file} ({self.output_file.stat().st_size} bytes)")
        return True


class WordFidelityVerifier(ExportTestRunner):
    """
    Specialized runner for Word Export fidelity checks.
    """
    def __init__(self, input_md_path, output_dir, export_func, report_file="fidelity_report.txt"):
        super().__init__(input_md_path, output_dir)
        self.export_func = export_func
        self.report_path = self.output_dir / report_file
        self.doc = None
        self.report_lines = []

    def run_export(self):
        self.output_file = self.output_dir / "test_output.docx"
        logger.info(f"Exporting to Word: {self.output_file}")
        
        docx_bytes = self.export_func(self.html_content)
        with open(self.output_file, "wb") as f:
            f.write(docx_bytes)

    def log_report(self, line):
        self.report_lines.append(line)
        # also print to console for visibility
        # print(line) 

    def inspect_docx(self):
        """Runs the suite of inspection tests on the generated DOCX."""
        from docx import Document
        from docx.oxml.ns import qn
        
        self.doc = Document(self.output_file)
        self.log_report(f"Inspecting {len(self.doc.tables)} tables in {self.output_file.name}...")
        
        self._check_alert_colors(qn)
        self._check_alert_icons()
        self._check_critic_markup()
        self._check_math_blocks(qn)
        self._check_abbreviations()
        self._check_details_transformation()
        self._check_emoji_rendering()
        
        # Save Report
        with open(self.report_path, "w", encoding="utf-8") as f:
            f.write("\n".join(self.report_lines))
        logger.info(f"Detailed Fidelity Report saved to: {self.report_path}")

    def _check_alert_colors(self, qn):
        self.log_report("-" * 40)
        self.log_report("TEST: Alert Colors")
        found_colored = 0
        for i, table in enumerate(self.doc.tables):
            if not table.rows or not table.columns: continue
            msg = table.cell(0,0).text.strip()[:15]
            
            tcPr = table.cell(0,0)._tc.get_or_add_tcPr()
            shd = tcPr.find(qn('w:shd'))
            fill = shd.get(qn('w:fill')) if shd is not None else "None"
            
            if shd is not None and fill not in ["None", "auto", "clear", "6366f1"]:
                self.log_report(f"Table {i} ({msg}...): Fill={fill} (PASS)")
                found_colored += 1
            else:
                self.log_report(f"Table {i} ({msg}...): Fill={fill}")
                
        self.log_report(f"Total Correctly Colored Alert Tables: {found_colored}")

    def _check_alert_icons(self):
        self.log_report("\n" + "-" * 40)
        self.log_report("TEST: Alert Icons (Segoe UI Emoji check)")
        icons = ["ℹ", "💡", "📣", "⚠️", "🛑", "⚡"]
        for table in self.doc.tables:
            if not table.rows: continue
            for p in table.cell(0,0).paragraphs:
                for run in p.runs:
                    if any(char in run.text for char in icons):
                        font = run.font.name
                        self.log_report(f"Icon Run '{run.text.strip()}' Font: {font}")

    def _check_critic_markup(self):
        self.log_report("\n" + "-" * 40)
        self.log_report("TEST: CriticMarkup (Colors/Styles)")
        for p in self.doc.paragraphs:
            if "Added" in p.text:
                for run in p.runs:
                    if "Added" in run.text and run.font.color:
                        self.log_report(f"Inserted '{run.text}': Color={run.font.color.rgb}")
            if "Deleted" in p.text:
                 for run in p.runs:
                    if "Deleted" in run.text:
                         self.log_report(f"Deleted '{run.text}': Strike={run.font.strike} Color={run.font.color.rgb if run.font.color else 'None'}")

    def _check_math_blocks(self, qn):
        self.log_report("\n" + "-" * 40)
        self.log_report("TEST: Math Blocks (Shading)")
        for p in self.doc.paragraphs:
            if "frac" in p.text: # Heuristic for the math example
                 shd = p._p.get_or_add_pPr().find(qn('w:shd'))
                 fill = shd.get(qn('w:fill')) if shd is not None else "None"
                 self.log_report(f"Math Block found: Fill={fill}")

    def _check_abbreviations(self):
        self.log_report("\n" + "-" * 40)
        self.log_report("TEST: Abbreviations")
        for p in self.doc.paragraphs:
             if "The HTML specification" in p.text:
                 if "Hyper Text Markup Language" in p.text:
                     self.log_report("FAIL: Abbreviation Expanded")
                 else:
                     self.log_report("PASS: Abbreviation Compact")

    def _check_details_transformation(self):
        self.log_report("\n" + "-" * 40)
        self.log_report("TEST: Details Transformation")
        found_details = False
        for p in self.doc.paragraphs:
            if "▶" in p.text and ("Click to Expand" in p.text or "Details" in p.text):
                found_details = True
                self.log_report(f"Found Details Header: '{p.text}'")
        
        if found_details:
             self.log_report("PASS: Details transformed to visual block.")
        else:
             self.log_report("FAIL: Details visual block not found (Icon missing?)")

    def _check_emoji_rendering(self):
         # This overlaps with check_alert_icons but targets the specific section
         self.log_report("\n" + "-" * 40)
         self.log_report("TEST: General Emoji Rendering")
         found_rocket = False
         for p in self.doc.paragraphs:
             for run in p.runs:
                 if "🚀" in run.text:
                     found_rocket = True
                     font = run.font.name
                     self.log_report(f"Found Rocket 🚀: Font={font}")
         
         if found_rocket:
             self.log_report("PASS: General Emoji found.")
         else:
             self.log_report("FAIL: General Emoji not found.")
