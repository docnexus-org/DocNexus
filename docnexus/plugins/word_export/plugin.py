import logging
import io
import shutil
import os
from pathlib import Path
import re
import urllib.parse
from flask import request, jsonify, send_file, current_app, Blueprint

# Note: Feature, FeatureType, FeatureState, PluginRegistry are INJECTED by the loader.
# Do not import them directly to avoid split-brain issues.

try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None

logger = logging.getLogger(__name__)

# Constants
MAX_EXPORT_HTML_SIZE = 50 * 1024 * 1024  # 50 MB

def add_bookmark(paragraph, bookmark_name):
    """Add a bookmark to a paragraph in a Word document."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    # Create bookmark start element
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), str(hash(bookmark_name) % 10000))
    bookmark_start.set(qn('w:name'), bookmark_name)
    
    # Create bookmark end element
    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), str(hash(bookmark_name) % 10000))
    
    # Insert bookmark
    # WARN: insert(0) puts it before pPr (properties), which creates INVALID XML that Word repairs by deleting.
    # We must append it to the end of the element (after pPr and existing runs).
    paragraph._element.append(bookmark_start)
    paragraph._element.append(bookmark_end)

# Imports for SafeHtmlToDocx and export_to_word
try:
    from htmldocx import HtmlToDocx
    from docx import Document
    from docx.shared import RGBColor, Pt, Inches
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR
    # re is imported globally
except ImportError as e:
    # Defer error handling to export_to_word if these are not available
    # This allows the module to load even if docx dependencies are missing
    HtmlToDocx = None
    Document = None
    RGBColor = None
    Pt = None
    OxmlElement = None
    qn = None
    WD_ALIGN_PARAGRAPH = None
    WD_COLOR = None
    re = None
    _word_export_import_error = e

# PDF Export Imports
try:
    from pdf2docx import Converter
except ImportError as e:
    Converter = None
    _pdf_export_import_error = e

class SafeHtmlToDocx(HtmlToDocx):
    """
    Subclass of HtmlToDocx to fix fragile color parsing that crashes on invalid hex/rgb strings.
    Overrides add_styles_to_run to add try/except blocks.
    """
    def add_styles_to_run(self, style):
        if 'color' in style:
            try:
                if 'rgb' in style['color']:
                    color = re.sub(r'[a-z()]+', '', style['color'])
                    parts = [x.strip() for x in color.split(',') if x.strip()]
                    if len(parts) >= 3:
                        colors = [int(p) for p in parts[:3]]
                        self.run.font.color.rgb = RGBColor(*colors)
                elif '#' in style['color']:
                    color = style['color'].strip().lstrip('#')
                    if len(color) == 3: color = ''.join([c*2 for c in color])
                    if len(color) >= 6:
                        colors = tuple(int(color[i:i+2], 16) for i in (0, 2, 4))
                        self.run.font.color.rgb = RGBColor(*colors)
            except Exception: pass
            
        if 'background-color' in style:
            try:
                bg = style['background-color'].lower()
                # CriticMarkup Mapping
                if '#ffff00' in bg: self.run.font.highlight_color = WD_COLOR.YELLOW
                elif '#008000' in bg: self.run.font.highlight_color = WD_COLOR.BRIGHT_GREEN
                elif '#ff0000' in bg: self.run.font.highlight_color = WD_COLOR.RED
                else: pass
            except Exception: pass
            
        # Text Decoration (Strike/Underline) from styles
        if 'text-decoration' in style:
            if 'line-through' in style['text-decoration']:
                self.run.font.strike = True
            if 'underline' in style['text-decoration']:
                self.run.font.underline = True

        # Font Family Support (For Emojis)
        if 'font-family' in style:
            try:
                # Log font found
                # logger.info(f"Font Family Found: {style['font-family']}")
                fonts = style['font-family'].split(',')
                if fonts:
                    primary_font = fonts[0].strip().replace("'", "").replace('"', "")
                    self.run.font.name = primary_font
            except Exception: pass

    def handle_starttag(self, tag, attrs):
        # Override to intercept Named Anchors for Bookmarks
        if tag == 'a':
            attrs_dict = dict(attrs)
            # Check for 'name' or 'id' that matches our Footnote pattern
            bookmark_name = attrs_dict.get('name') or attrs_dict.get('id')
            
            if bookmark_name and (bookmark_name.startswith('fn_') or bookmark_name.startswith('fnref_')):
                # We found a footnote anchor! Inject a Word Bookmark.
                # Note: We need a paragraph to attach to. htmldocx usually manages self.paragraph.
                if self.paragraph:
                    # logger.info(f"Injecting Word Bookmark: {bookmark_name}")
                    add_bookmark(self.paragraph, bookmark_name)
                    
        super().handle_starttag(tag, attrs)

    def add_styles_to_paragraph(self, style):
        # Override to support background-color (Shading) for Math Blocks
        super().add_styles_to_paragraph(style)
        
        if 'background-color' in style:
            try:
                color = style['background-color'].strip().lstrip('#')
                if len(color) == 3: color = ''.join([c*2 for c in color])
                if len(color) >= 6:
                    # Log finding color
                    # logger.info(f"Injecting Paragraph Shading: {color}")
                    pPr = self.paragraph._p.get_or_add_pPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:fill'), color)
                    
                    # Schema Order for pPr: ... pBdr, shd, tabs, spacing, ind, jc, rPr ...
                    successors = ['w:tabs', 'w:spacing', 'w:ind', 'w:jc', 'w:rPr']
                    target = None
                    for s in successors:
                        target = pPr.find(qn(s))
                        if target is not None:
                            break
                    
                    if target is not None:
                        pPr.insert_element_before(shd, target.tag)
                    else:
                        pPr.append(shd)
            except Exception as e:
                logger.error(f"Error injecting shading: {e}")


def transform_html_for_word(soup: BeautifulSoup):
    """
    Transforms HTML elements into Word-friendly structures.
    Modifies the soup in-place.
    """
    # 1. Transform Tabs (.tabbed-set) -> Vertical Headings + Content
    # Structure: .tabbed-set > input, label, .tabbed-content
    for tab_set in soup.find_all(class_='tabbed-set'):
        # Create a container for the flattened content
        flattened_div = soup.new_tag('div')
        
        # Iterate over labels and corresponding content
        labels = tab_set.find_all('label')
        contents = tab_set.find_all(class_='tabbed-content')
        
        for i, label in enumerate(labels):
            if i < len(contents):
                # Create Heading from Label
                h4 = soup.new_tag('h4')
                h4.string = label.get_text(strip=True)
                h4['style'] = "margin-top: 12px; margin-bottom: 4px; color: #4b5563;"
                flattened_div.append(h4)
                
                # Append Content directly
                content = contents[i]
                # Remove class to prevent CSS interference if any
                if 'class' in content.attrs:
                    del content['class']
                content['style'] = "margin-left: 8px; margin-bottom: 12px;"
                flattened_div.append(content)
        
        # Replace the complex tab set with the flattened div
        tab_set.replace_with(flattened_div)

    # 2. Transform Collapsible Details (details) -> Bold Summary + Content
    for details in soup.find_all('details'):
        summary = details.find('summary')
        if summary:
            # Create a bold paragraph for the summary
            p = soup.new_tag('p')
            b = soup.new_tag('b')
            b.string = f"▶ {summary.get_text(strip=True)}"
            p.append(b)
            p['style'] = "margin-top: 8px; margin-bottom: 4px;"
            
            # Insert summary P before the details tag
            details.insert_before(p)
            
            # Unwrap the details tag (keeping children, removing details wrapper)
            # The summary tag is still there, need to remove it
            summary.decompose()
            details.unwrap()

    # 3. Transform GitHub Alerts (.admonition) -> Single-Cell Tables
    # htmldocx doesn't support complex borders/backgrounds on divs well.
    alert_themes = {
        'note':      {'border': '#0969da', 'bg': '#e6f6ff', 'icon': 'ℹ️'},  # Blue
        'tip':       {'border': '#1a7f37', 'bg': '#dafbe1', 'icon': '💡'},  # Green
        'important': {'border': '#8250df', 'bg': '#f3e6ff', 'icon': '📣'},  # Purple
        'warning':   {'border': '#bf8700', 'bg': '#fff8c5', 'icon': '⚠️'},  # Amber (Fixed Hex)
        'caution':   {'border': '#d1242f', 'bg': '#ffebe9', 'icon': '🛑'},  # Red
        'danger':    {'border': '#d1242f', 'bg': '#ffebe9', 'icon': '⚡'}   # Red/Danger
    }
    
    for admonition in soup.find_all(class_='admonition'):
        # Determine type/color
        classes = admonition.get('class', [])
        color = '#0969da' # Default Blue
        bg_color = '#e6f6ff' # Default Light Blue
        icon = 'ℹ️'
        alert_type = 'NOTE'
        
        for cls, theme in alert_themes.items():
            if cls in classes:
                color = theme['border']
                bg_color = theme['bg']
                icon = theme['icon']
                alert_type = cls.upper()
                break
        
        table = soup.new_tag('table')
        # Add marker class so we can skip global table styling later
        table['class'] = 'docnexus-alert-table'
        # Use full border as htmldocx/Word support for partial borders is flaky
        table['style'] = f"border-collapse: collapse; width: 100%; border: 2px solid {color}; background-color: {bg_color};"
        tr = soup.new_tag('tr')
        td = soup.new_tag('td')
        td['style'] = f"padding: 8px; background-color: {bg_color};" # Apply BG to TD as well for safety
        
        # Extract Title
        title = admonition.find(class_='admonition-title')
        if title:
            # Create a bold paragraph for the title
            title_p = soup.new_tag('p')
            title_b = soup.new_tag('b')
            title_span = soup.new_tag('span')
            
            title_text = title.get_text(strip=True) or alert_type
            # User reported extra spaces. Removing explicit space.
            title_span.string = f"{icon}\u00A0{title_text}" # Using non-breaking space for consistent small gap
            
            # Force Emoji Font for color rendering (applied to span, as b tag styles might be ignored)
            # Simplify font string to avoid parsing issues with quotes/commas
            title_span['style'] = f"color: {color}; font-family: Segoe UI Emoji;"
            
            title_b.append(title_span)
            title_p.append(title_b)
            td.append(title_p)
            title.decompose()
        
        # Move remaining content to TD
        # We need to copy children one by one to avoid issues while modifying the tree
            
        # Move content to cell
        content_div = soup.new_tag('div')
        # Move siblings of title to content_div? No, GitHub Alerts flat structure is tricky.
        # Usually alert content follows the Title blockquote or similar.
        # But here we are assuming the whole 'alert' div content is what we want.
        # We already decomposed title, so the rest is content.
        for child in list(admonition.contents):
            content_div.append(child)
            
        td.append(content_div)
        tr.append(td)
        table.append(tr)
        admonition.replace_with(table)

    # 4. Transform Details (Collapsible) -> Styled Block
    # <details><summary>Title</summary>Content</details>
    for details in soup.find_all('details'):
        summary = details.find('summary')
        
        # Create container div
        container = soup.new_tag('div')
        container['style'] = "border: 1px solid #cccccc; padding: 10px; margin: 10px 0; background-color: #fafafa;"
        
        # Handle Title
        title_p = soup.new_tag('p')
        title_b = soup.new_tag('b')
        if summary:
            # Add icon and text
            title_b.string = f"▶ {summary.get_text(strip=True)}"
            summary.decompose() # Remove summary from details content
        else:
            title_b.string = "▶ Details"
        
        title_p.append(title_b)
        container.append(title_p)
        
        # Handle Content (Remaining children of details)
        content_div = soup.new_tag('div')
        content_div['style'] = "margin-top: 5px; margin-left: 15px;"
        
        # Move remaining contents
        for child in list(details.contents):
             content_div.append(child)
             
        container.append(content_div)
        details.replace_with(container)

    # 5. Transform Emojis (Wrap in Font Span)
    # Walk text nodes to find emojis and wrap them
    import re
    # Regex for common emojis (including supplementary pairs)
    # Simplified regex for the requested ones + ranges
    emoji_pattern = re.compile(r'[\U0001F600-\U0001F64F\U0001F300-\U0001F5FF\U0001F680-\U0001F6FF\U0001F700-\U0001F77F\U0001F780-\U0001F7FF\U0001F800-\U0001F8FF\U0001F900-\U0001F9FF\U0001FA00-\U0001FA6F\u2600-\u26FF\u2700-\u27BF]')
    
    # We iterate a list of text nodes to modify them safely
    for text_node in soup.find_all(string=True):
        if text_node.parent and text_node.parent.name in ['script', 'style']:
            continue
            
        if emoji_pattern.search(text_node):
            # If emoji found, we need to split and wrap
            new_content = []
            last_idx = 0
            for match in emoji_pattern.finditer(text_node):
                start, end = match.span()
                # Text before
                if start > last_idx:
                    new_content.append(soup.new_string(text_node[last_idx:start]))
                
                # Emoji Wrapped
                emoji_span = soup.new_tag('span')
                # Append VS16 (\ufe0f) to force Emoji Presentation
                emoji_char = text_node[start:end]
                if not emoji_char.endswith('\ufe0f'):
                     emoji_char += '\ufe0f'
                
                emoji_span['style'] = "font-family: 'Segoe UI Emoji', sans-serif;"
                emoji_span.string = emoji_char
                new_content.append(emoji_span)
                
                last_idx = end
            
            # Text after
            if last_idx < len(text_node):
                new_content.append(soup.new_string(text_node[last_idx:]))
            
            # Replace text node with new structure
            # We use a span as a container if parent allows, or insert sibling
            # replace_with allows passing multiple arguments!
            text_node.replace_with(*new_content)

    # 6. Transform Task Lists -> Text [x] / [ ]
    for checkbox in soup.find_all('input', {'type': 'checkbox'}):
        is_checked = checkbox.has_attr('checked')
        replacement = soup.new_tag('span')
        replacement.string = "[x] " if is_checked else "[ ] "
        replacement['style'] = "font-family: monospace;"
        checkbox.replace_with(replacement)
        
    # 5. Transform CriticMarkup (Legacy numbering, keeping order)
    # Highlight
    for mark in soup.find_all('mark'):
        mark.name = 'span'
        # Yellow background for Word (will be mapped in add_styles_to_run override)
        mark['style'] = "background-color: #ffff00;"
        
    # Insert (Underline)
    for ins in soup.find_all('ins'):
        ins.name = 'span'
        ins['style'] = "color: #008000; text-decoration: underline;" # Green text + underline
        
    # Delete (Strikethrough)
    for delete in soup.find_all('del'):
        delete.name = 'span'
        delete['style'] = "color: #ff0000; text-decoration: line-through;" # Red text + strike

    # 6. Transform Definition Lists (dl, dt, dd) -> Bold + Indent
    for dl in soup.find_all('dl'):
        # We unwrap the dl, and style dt/dd
        for dt in dl.find_all('dt'):
            dt.name = 'p'
            b = soup.new_tag('b')
            b.string = dt.get_text(strip=True)
            dt.string = ''
            dt.append(b)
            dt['style'] = "margin-top: 8px; margin-bottom: 2px;"
            
        for dd in dl.find_all('dd'):
            dd.name = 'p'
            dd['style'] = "margin-left: 20px; margin-bottom: 8px;"
            
        dl.unwrap()

    # 7. Transform Math currently handled by Image Generation Logic later
    # We purposefully skip it here to preserve the nodes for the Image Converter.

    # 8. Transform Footnotes (Header + Styles + Hyperlinks)
    footnote_div = soup.find('div', class_='footnote')
    if footnote_div:
        # A. Add Header
        header = soup.new_tag('h3')
        header.string = "Footnotes"
        header['style'] = "margin-top: 24px; margin-bottom: 12px; border-bottom: 1px solid #cccccc; padding-bottom: 4px; text-align: left;"
        footnote_div.insert_before(header)
        
        # B. Style Container
        footnote_div['style'] = "font-size: 10pt; color: #4b5563; text-align: left;"
        
        # C. Fix Hyperlinks & Layout
        # Word Bookmarks require <a name="...">. 
        # Crucial: htmldocx often ignores empty anchors <a name="foo"></a>.
        # Fix: Put a Zero-Width Space (&#8203;) inside target anchors.
        # Fix: Merge 'name' into the existing link for references.
        
        # C1. Targets (The Footnote Definitions)
        for li in footnote_div.find_all('li', id=True):
            raw_id = li['id']
            # We move the ID to the anchor to ensure htmldocx sees it as a bookmark target
            del li['id'] 
            safe_id = raw_id.replace(':', '_')
            
            # 1. Inject Anchor with Content (ZWS)
            # Use both name and id to be safe
            a_target = soup.new_tag('a')
            a_target['name'] = safe_id
            a_target['id'] = safe_id
            # Fix KeyError: 'href' in htmldocx
            # htmldocx assumes any <a> with content has an href.
            a_target['href'] = '#' 
            a_target.string = "\u200b" # Zero Width Space
            li.insert(0, a_target)
            
            # 2. Update Backlink Href
            backlink = li.find('a', class_='footnote-backref')
            if backlink and backlink.get('href'):
                 raw_href = backlink['href']
                 safe_href = raw_href.replace(':', '_')
                 backlink['href'] = safe_href
            
            # 3. Unwrap Paragraphs
            for p in li.find_all('p'):
                p.unwrap()
                
            # 4. Enforce Left Alignment
            li['style'] = "text-align: left; margin-bottom: 4px;"

        # C2. Sources (The References in Text)
        for sup in soup.find_all('sup', id=re.compile(r'^fnref')):
            raw_id = sup['id']
            del sup['id']
            safe_id = raw_id.replace(':', '_')
            
            # 1. Find existing link
            a_link = sup.find('a', class_='footnote-ref')
            if a_link:
                # Merge bookmark identity into the link itself
                a_link['name'] = safe_id
                a_link['id'] = safe_id
                
                # Fix Forward Link Href
                if a_link.get('href'):
                     raw_href = a_link['href']
                     safe_href = raw_href.replace(':', '_')
                     a_link['href'] = safe_href
            else:
                 # Fallback if weird structure: inject anchor with ZWS?
                 # Should not happen with standard markdown
                 pass

def export_to_word(html_content: str) -> bytes:
    """
    Exports HTML content to a Word (.docx) file byte stream.
    """
    if HtmlToDocx is None:
        logger.error(f"Failed to import Word export dependencies: {_word_export_import_error}")
        raise RuntimeError("Word export dependencies (htmldocx, python-docx) not installed.")

    # Size Check
    html_size = len(html_content.encode('utf-8'))
    if html_size > MAX_EXPORT_HTML_SIZE:
        raise ValueError(f"Content too large ({html_size/1024/1024:.2f} MB). Max {MAX_EXPORT_HTML_SIZE/1024/1024} MB.")

    logger.info(f"WordExport: Generating document from {html_size} bytes of HTML...")

    # Pre-process HTML with BeautifulSoup
    try:
        soup = BeautifulSoup(html_content, 'lxml')
    except:
        soup = BeautifulSoup(html_content, 'html.parser')
    
    # Cleaning (Scripts, Styles, Nav)
    # CRITICAL: Do NOT delete math scripts yet, we need them for extraction!
    for tag in soup.find_all(['script', 'style', 'nav']):
        if tag.name == 'script' and tag.get('type') and 'math/tex' in tag.get('type'):
            continue
        tag.decompose()
        
    # LOGGING: Inspect HTML Structure entering Word Transform (CRITICAL DEBUG)
    # logger.info(f"WordExport: HTML Head Snippet: {soup.prettify()[:2000]}")
    
    # Transform Complex HTML for Word Compatibility
    # (Tabs, Alerts, Details, Math, etc.)
    transform_html_for_word(soup)
    
    # Main Content Extraction
    # We want to include the Table of Contents (.toc-container) AND the Markdown Content (.markdown-content)
    # The frontend wraps both in #documentContent div (Line 729 view.html)
    # But usually sending full <html>.
    
    container = soup.find(id='documentContent')
    selected_content = []
    
    if container:
        # Extract TOC if present
        toc = container.find(class_='toc-container')
        if toc:
             # Style TOC for Word
            toc_header = toc.find(class_='toc-header')
            if toc_header:
                toc_header.name = 'h2' # Make it a standard header for Word
                toc_header['style'] = 'font-size: 14pt; color: #4b5563; margin-top: 0;'
            
            selected_content.append(toc)
            
            # Robust Page Break: Inject a unique marker we can find and replace with a REAL Word Break later
            # CSS page-break-after is unreliable in htmldocx
            pb_marker = soup.new_tag('p')
            pb_marker.string = "<<<DOCNEXUS_PAGE_BREAK>>>"
            selected_content.append(pb_marker)
            
        # Extract Markdown Content
        md_content = container.find(class_='markdown-content')
        if md_content:
             selected_content.append(md_content)
    else:
        # Fallback to old behavior if ID not found
        md_content = soup.find(class_='markdown-content')
        if md_content:
             selected_content.append(md_content)

    if selected_content:
        logger.info(f"WordExport: Analyzing {len(selected_content)} content parts for tables.")

        # Style Tables for Word (Apply to all tables in selected content)
        for part in selected_content:
            for table in part.find_all('table'):
                # 1. Logging & Classification
                classes = table.get('class', [])
                # Normalize class attribute to list (bs4 can return str or list)
                class_list = classes if isinstance(classes, list) else classes.split() if isinstance(classes, str) else []
                
                is_alert = 'docnexus-alert-table' in class_list
                
                # 2. Logic
                if is_alert:
                    # logger.info(f"WordExport: Skipping Global Style for Alert Table. Classes={class_list}")
                    continue
                
                # Standard Table Styling
                # logger.debug(f"WordExport: Applying Global Style to Standard Table. Classes={class_list}")
                table['style'] = 'border-collapse: collapse; width: 100%; border: 2px solid #6366f1; margin-bottom: 20px;'
                table['border'] = '1'
                
                # Thead check
                thead = table.find('thead')
                if not thead:
                    first_row = table.find('tr')
                    if first_row and first_row.find('th'):
                        thead = soup.new_tag('thead')
                        first_row.extract()
                        thead.append(first_row)
                        table.insert(0, thead)
                
                # Colors and Styles injection
                for th in table.find_all('th'):
                    th['bgcolor'] = '#6366f1'
                    th['style'] = 'background-color: #6366f1 !important; color: #ffffff !important;'
                
                for td in table.find_all('td'):
                    td['style'] = 'padding: 8px; border: 1px solid #e5e7eb;'

        # Combine content
        combined_html = "".join([str(tag) for tag in selected_content])
        clean_html = f'<html><head><meta charset="utf-8"></head><body>{combined_html}</body></html>'
        
        # Capture main_content for booking logic later (use md_content reference)
        main_content = md_content if 'md_content' in locals() and md_content else None
    else:
        logger.warning("WordExport: No 'selected_content' found to style! using absolute fallback.")
        # Absolute fallback
        clean_html = f'<html><body>{soup.body.decode_contents() if soup.body else str(soup)}</body></html>'
        main_content = None

    # Pre-process HTML to resolve/fetch images (crucial for stability)
    # This prevents htmldocx from crashing on network errors or missing files.
    soup = BeautifulSoup(clean_html, 'html.parser')
    
    # We need a temp dir for downloaded images that persists during conversion
    import tempfile
    import urllib.request
    from urllib.parse import urlparse
    import shutil
    import re # Ensure re is available
    
    def parse_tex_to_html(soup_factory, tex_str):
        """
        Converts simple TeX (subscripts, superscripts, basic symbols) to an HTML span.
        Handles: x^2, x_i, x_{i+1}, \alpha -> alpha
        """
        container = soup_factory.new_tag('span')
        container['style'] = "font-family: 'Cambria Math', 'Times New Roman', serif;"
        
        cursor = 0
        n = len(tex_str)
        
        while cursor < n:
            char = tex_str[cursor]
            
            if char in ('^', '_'):
                tag_name = 'sup' if char == '^' else 'sub'
                cursor += 1
                content = ""
                
                # Check for Group { }
                if cursor < n and tex_str[cursor] == '{':
                    cursor += 1
                    nesting = 1
                    start_grp = cursor
                    while cursor < n and nesting > 0:
                        if tex_str[cursor] == '{': nesting += 1
                        elif tex_str[cursor] == '}': nesting -= 1
                        if nesting > 0: cursor += 1
                        
                    content = tex_str[start_grp:cursor]
                    cursor += 1 # Skip closing }
                elif cursor < n:
                    if tex_str[cursor] == '\\':
                         # Extract macro
                         macro_start = cursor
                         cursor += 1
                         while cursor < n and tex_str[cursor].isalpha():
                             cursor += 1
                         content = tex_str[macro_start:cursor]
                    else:
                        content = tex_str[cursor]
                        cursor += 1
                        
                # Create Tag
                elem = soup_factory.new_tag(tag_name)
                # Recursively parse content? Or just strip braces?
                # Simple recursion for nested superscripts
                if '^' in content or '_' in content:
                    nested_span = parse_tex_to_html(soup_factory, content)
                    # Unwrap span into elem
                    for child in list(nested_span.contents):
                        elem.append(child)
                else:
                    key = content.replace('\\', '').strip()
                    # Basic Mapper (could be expanded)
                    symbols = {
                        'alpha': 'α', 'beta': 'β', 'gamma': 'γ', 'theta': 'θ', 'pi': 'π', 
                        'sigma': 'σ', 'omega': 'ω', 'Delta': 'Δ', 'mu': 'μ', 'lambda': 'λ',
                        'infty': '∞', 'rightarrow': '→', 'leftarrow': '←', 'approx': '≈',
                        'neq': '≠', 'le': '≤', 'ge': '≥', 'times': '×', 'cdot': '·'
                    }
                    elem.string = symbols.get(key, content.replace('\\', ''))
                
                container.append(elem)
                
            elif char == '{' or char == '}':
                 cursor += 1
            elif char == '\\':
                # Handle text macros in main flow
                macro_start = cursor
                cursor += 1
                while cursor < n and tex_str[cursor].isalpha():
                    cursor += 1
                macro = tex_str[macro_start+1:cursor]
                
                symbols = {
                    'alpha': 'α', 'beta': 'β', 'gamma': 'γ', 'theta': 'θ', 'pi': 'π',
                    'sigma': 'σ', 'omega': 'ω', 'Delta': 'Δ', 'mu': 'μ', 'lambda': 'λ',
                    'infty': '∞', 'rightarrow': '→', 'leftarrow': '←', 'approx': '≈',
                     'neq': '≠', 'le': '≤', 'ge': '≥', 'times': '×', 'cdot': '·', 'frac': ''
                }
                
                if macro == 'frac':
                     # simple frac ignore?
                     pass
                else:
                    container.append(soup_factory.new_string(symbols.get(macro, ""))) # Only append if mapped, otherwise skip backslash
            else:
                # Regular text
                # Gather contiguous text
                text_buffer = ""
                while cursor < n and tex_str[cursor] not in ('^', '_', '{', '}', '\\'):
                    text_buffer += tex_str[cursor]
                    cursor += 1
                container.append(soup_factory.new_string(text_buffer))
                
        return container
    
    # Create a temporary directory for this export session
    with tempfile.TemporaryDirectory() as temp_img_dir:
        # 4. Transform Math (KaTeX/MathJax) -> Image (CodeCogs)
        # Target: .katex-mathml annotation[encoding="application/x-tex"] or <script type="math/tex">
        
        # --- PRE-CLEANUP: Remove MathJax Previews ---
        # These often sit adjacent to the script tag and just clutter the DOM.
        for junk in soup.find_all(class_=['MathJax_Preview', 'katex-html']):
            # Verify valid parent to avoid double-deletion errors
            if junk.parent:
                junk.decompose()
                
        # Collect candidates
        # --- PASS 1: Specific Math Elements (Scripts & KaTeX spans) ---
        # We target the leaf nodes first to ensure granularity (e.g. multiple formulas in one line/container)
        
        specific_candidates = []
        specific_candidates.extend(soup.find_all('script', type=re.compile(r'math/tex')))
        # Target .katex but ignore those inside other .katex (nested)
        for k in soup.find_all(class_='katex'):
             if not k.find_parent(class_='katex'):
                 specific_candidates.append(k)
        
        processed_math_ids = set()
        
        # logger.info(f"WordExport: Found {len(specific_candidates)} specific math candidates (Scripts/KaTeX).")
        
        for i, target_node in enumerate(specific_candidates):
            if id(target_node) in processed_math_ids:
                continue
            
            # Extract TeX
            tex = ""
            is_display = False
            
            script_child = target_node if target_node.name == 'script' else target_node.find('script', type=re.compile(r'math/tex'))
            annotation_child = target_node.find('annotation', attrs={'encoding': 'application/x-tex'})
            if not annotation_child:
                annotation_child = target_node.find('annotation')
                
            if script_child:
                tex = script_child.get_text()
                is_display = 'mode=display' in script_child.get('type', '')
            elif annotation_child:
                tex = annotation_child.get_text().strip()
                is_display = (target_node.name == 'div') or \
                             ('display' in (target_node.get('class') or [])) or \
                             ('katex-display' in target_node.decode_contents()) or \
                             (target_node.find_parent(class_='katex-display') is not None)
            
            if tex and tex.strip():
                tex = tex.strip()
                processed_math_ids.add(id(target_node))
                
                if is_display:
                    # BLOCK MATH: Use Image (CodeCogs) for full fidelity
                    base_url = "https://latex.codecogs.com/png.image"
                    # Add \displaystyle to ensure block rendering if not present
                    if not tex.startswith(r'\displaystyle'):
                        tex = r'\displaystyle ' + tex
                        
                    # Encoded params
                    params = {
                        'dpi': 300,
                        'bg': 'white'
                    }
                    # CodeCogs expects strict URL encoding of the latex
                    encoded_tex = urllib.parse.quote(tex)
                    # Construct full URL manually: base + \dpi{300} + tex
                    # CodeCogs Format: https://latex.codecogs.com/png.image?\dpi{300}\bg{white}formula
                    full_url = f"{base_url}?\\dpi{{300}}{encoded_tex}"
                    
                    # Create Image Tag
                    img = soup.new_tag('img')
                    img['src'] = full_url
                    img['alt'] = tex
                    img['class'] = 'math-block-image'
                    img['style'] = "display: block; margin: 10px auto; max-width: 90%;"
                    
                    # Replace
                    target_node.replace_with(img)
                else:
                    # INLINE MATH: Convert to Valid Word Text (Subscripts/Superscripts)
                    # Using CodeCogs for inline images often creates alignment hell in Word.
                    # We will try a text-based approximation buffer.
                    
                    # Parse TeX (Simple Superscript/Subscript)
                    try:
                        span = parse_tex_to_html(soup, tex)
                        target_node.replace_with(span)
                    except Exception as e:
                        # Fallback
                        fallback = soup.new_tag('span')
                        fallback.string = f"[{tex}]"
                        fallback['style'] = "color: #555; font-family: monospace;"
                        target_node.replace_with(fallback)

        # --- PASS 2: Download Remote Images ---
        # Iterate all images, download to temp dir, and point src to local file path
        # HtmlToDocx needs local paths to embed correctly usually?
        # No, HtmlToDocx handles URLs but fails if it can't fetch.
        # We pre-fetch to ensure robustness and valid headers (User-Agent).
        
        for img in soup.find_all('img'):
            src = img.get('src')
            if not src: continue
            
            # Skip valid local file paths (if any)
            if os.path.exists(src): continue
            
            # Skip Base64
            if src.startswith('data:'): continue
            
            try:
                # Resolve relative URLS (assume against localhost:5000 if needed, or skip)
                if src.startswith('/'):
                    # Local server asset?
                    # We can't easily fetch from "myself" inside the plugin without full URL.
                    # Better to skip or assume it's reachable via http://localhost:PORT
                    # For now, let's try to ignore relative unless we know the domain.
                    pass
                else:
                    # Remote URL (CodeCogs, etc.)
                   # Generate Temp Filename
                   parsed = urlparse(src)
                   ext = os.path.splitext(parsed.path)[1]
                   if not ext: ext = '.png'
                   
                   temp_filename = f"img_{hash(src)}{ext}"
                   temp_path = os.path.join(temp_img_dir, temp_filename)
                   
                   # Download with requests
                   headers = {'User-Agent': 'DocNexus/1.0'}
                   import requests
                   
                   # Timeout to prevent hanging
                   r = requests.get(src, headers=headers, timeout=5)
                   if r.status_code == 200:
                       with open(temp_path, 'wb') as f:
                           f.write(r.content)
                       
                       # Update SRC to absolute local path
                       img['src'] = temp_path
                   else:
                       logger.warning(f"Failed to download image: {src} (Status {r.status_code})")
            except Exception as e:
                logger.warning(f"Error pre-fetching image {src}: {e}")
                
        # 5. Convert to Docx
        docx_buffer = io.BytesIO()
        document = Document()
        new_parser = SafeHtmlToDocx()
        
        # Capture generated content
        clean_html_final = str(soup)
        
        # DEBUG: Save sanitized HTML to file for inspection
        # with open("debug_export_sanitized.html", "w", encoding="utf-8") as f:
        #    f.write(clean_html_final)
        
        new_parser.add_html_to_document(clean_html_final, document)
        
        # 6. Post-Process Docx (Page Breaks)
        # Iterate paragraphs and look for our marker
        for p in document.paragraphs:
            if "<<<DOCNEXUS_PAGE_BREAK>>>" in p.text:
                # Clear text
                p.text = ""
                # Add run with break
                run = p.add_run()
                run.add_break() # Default is PAGE break
                
        document.save(docx_buffer)
        docx_buffer.seek(0)
        
        return docx_buffer.getvalue()

def export_pdf_to_docx(pdf_abs_path: str) -> bytes:
    """
    Converts a PDF file to DOCX using pdf2docx.
    """
    if Converter is None:
        logger.error(f"Failed to import PDF conversion dependencies: {_pdf_export_import_error}")
        raise RuntimeError("PDF to DOCX dependencies (pdf2docx) not installed. Please run `pip install pdf2docx`.")

    if not os.path.exists(pdf_abs_path):
        raise FileNotFoundError(f"PDF file not found: {pdf_abs_path}")

    logger.info(f"PDFExport: Converting {pdf_abs_path} to DOCX...")
    
    # Use temp file for output
    import tempfile
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_docx:
        temp_docx_path = tmp_docx.name
    
    try:
        cv = Converter(pdf_abs_path)
        cv.convert(temp_docx_path, start=0, end=None)
        cv.close()
        
        with open(temp_docx_path, 'rb') as f:
            docx_bytes = f.read()
            
        return docx_bytes
    except Exception as e:
        logger.error(f"PDF Conversion failed: {e}")
        raise
    finally:
        # Cleanup
        if os.path.exists(temp_docx_path):
            try:
                os.remove(temp_docx_path)
            except: pass


# -------------------------------------------------------------------------
# Flask Blueprint Routes
# -------------------------------------------------------------------------
# Blueprint is created in __init__ usually, but here we can't assume structure.
# We will check if `get_blueprint` is expected or we return a blueprint object.
# Standard DocNexus plugins return a function `register_blueprint(app)`.
# Or expect `plugin_bp` to be exposed.

from flask import Blueprint

blueprint = Blueprint('word_export', __name__)

@blueprint.route('/api/export/docx', methods=['POST'])
def export_docx_route():
    try:
        data = request.json
        html_content = data.get('content')
        filename = data.get('filename', 'document.docx')
        
        if not html_content:
            return jsonify({'error': 'No content provided'}), 400
            
        docx_bytes = export_to_word(html_content)
        
        return send_file(
            io.BytesIO(docx_bytes),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename.replace('.md', '.docx').replace('.html', '.docx')
        )
    except Exception as e:
        logger.error(f"Export failed: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500

@blueprint.route('/api/export/pdf-to-docx', methods=['POST'])
def export_pdf_to_docx_route():
    try:
        data = request.json
        file_path_rel = data.get('filePath')
        
        if not file_path_rel:
            return jsonify({'error': 'No file path provided'}), 400
            
        # Security: Resolve Path
        workspace = current_app.config.get('WORKSPACE_PATH', '')
        # Remove leading slash if present to join correctly
        clean_rel = file_path_rel.lstrip('/\\')
        abs_path = os.path.join(workspace, clean_rel)
        
        if not os.path.exists(abs_path):
             return jsonify({'error': f'File not found: {file_path_rel}'}), 404
             
        docx_bytes = export_pdf_to_docx(abs_path)
        
        filename = os.path.basename(abs_path).replace('.pdf', '.docx')
        
        return send_file(
            io.BytesIO(docx_bytes),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        logger.error(f"PDF to DOCX failed: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500

import traceback

# Plugin Entry Point
# Note: Feature, FeatureType, FeatureState are injected into this module's globals by the loader.

def get_features():
    return [
        Feature(
            name="Word Export",
            handler=export_to_word,
            feature_type=FeatureType.EXPORT_HANDLER,
            state=FeatureState.STANDARD,
            meta={
                "extension": "docx", # Standardize meta key for export handlers
                "label": "Word Document (.docx)",
                "content_type": "text/html" 
            }
        )
    ]


PLUGIN_METADATA = {
    'name': 'Word Export',
    'description': 'Converts documents and PDFs to professional Microsoft Word (.docx) format.',
    'category': 'export',
    'icon': 'fa-file-word',
    'preinstalled': True # Ensure it loads by default if possible, or matches user state
}
